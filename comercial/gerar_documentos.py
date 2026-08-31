#!/usr/bin/env python3
"""Gera PDF e DOCX comerciais do SGE Municipal (agosto/2026)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT = Path(__file__).resolve().parent
DATA = "30 de agosto de 2026"
REF = "PC-SGE-001/2026"
VALIDADE = "45 (quarenta e cinco) dias corridos"

INK = colors.HexColor("#1B2433")
BRAND = colors.HexColor("#1E4A7A")
BRAND_DARK = colors.HexColor("#152C4A")
LINE = colors.HexColor("#D5DCE6")
MUTED = colors.HexColor("#5A6573")
SURFACE = colors.HexColor("#F4F6F9")
OK = colors.HexColor("#1B6B4A")
WARN = colors.HexColor("#8A5A12")
WHITE = colors.white

DOCX_INK = RGBColor(0x1B, 0x24, 0x33)
DOCX_BRAND = RGBColor(0x1E, 0x4A, 0x7A)
DOCX_MUTED = RGBColor(0x5A, 0x65, 0x73)
DOCX_OK = RGBColor(0x1B, 0x6B, 0x4A)
DOCX_WARN = RGBColor(0x8A, 0x5A, 0x12)


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------


def _styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    s: dict[str, ParagraphStyle] = {}
    s["cover_kicker"] = ParagraphStyle(
        "cover_kicker",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=9,
        textColor=WHITE,
        letterSpacing=1.2,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    s["cover_title"] = ParagraphStyle(
        "cover_title",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=22,
        leading=26,
        textColor=WHITE,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    s["cover_sub"] = ParagraphStyle(
        "cover_sub",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#D7E2F0"),
        alignment=TA_LEFT,
    )
    s["h1"] = ParagraphStyle(
        "h1",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=14,
        leading=18,
        textColor=BRAND_DARK,
        spaceBefore=16,
        spaceAfter=8,
    )
    s["h2"] = ParagraphStyle(
        "h2",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=12,
        leading=16,
        textColor=BRAND,
        spaceBefore=12,
        spaceAfter=6,
    )
    s["body"] = ParagraphStyle(
        "body",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=10,
        leading=14,
        textColor=INK,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    s["body_left"] = ParagraphStyle(
        "body_left",
        parent=s["body"],
        alignment=TA_LEFT,
    )
    s["small"] = ParagraphStyle(
        "small",
        parent=base["Normal"],
        fontName="Times-Italic",
        fontSize=8.5,
        leading=11,
        textColor=MUTED,
        spaceAfter=8,
    )
    s["cell"] = ParagraphStyle(
        "cell",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=8,
        leading=11,
        textColor=INK,
    )
    s["cell_h"] = ParagraphStyle(
        "cell_h",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=8,
        leading=11,
        textColor=WHITE,
    )
    s["bullet"] = ParagraphStyle(
        "bullet",
        parent=s["body_left"],
        leftIndent=12,
        spaceAfter=3,
    )
    s["meta"] = ParagraphStyle(
        "meta",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=9,
        leading=12,
        textColor=INK,
    )
    s["center"] = ParagraphStyle(
        "center",
        parent=s["body"],
        alignment=TA_CENTER,
    )
    s["footer"] = ParagraphStyle(
        "footer",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=8,
        textColor=MUTED,
        alignment=TA_LEFT,
    )
    s["verdict"] = ParagraphStyle(
        "verdict",
        parent=base["Normal"],
        fontName="Times-Bold",
        fontSize=10.5,
        leading=14,
        textColor=OK,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    s["warn"] = ParagraphStyle(
        "warn",
        parent=base["Normal"],
        fontName="Times-Roman",
        fontSize=10,
        leading=14,
        textColor=WARN,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    return s


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("\n", "<br/>"), style)


def _table(headers: list[str], rows: list[list[str]], col_widths: list[float], s: dict) -> Table:
    head = [_p(h, s["cell_h"]) for h in headers]
    data = [head]
    for row in rows:
        data.append([_p(c, s["cell"]) for c in row])
    t = Table(data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_DARK),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("BACKGROUND", (0, 1), (-1, -1), WHITE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("GRID", (0, 0), (-1, -1), 0.4, LINE),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            style_cmds.append(("BACKGROUND", (0, i), (-1, i), SURFACE))
    t.setStyle(TableStyle(style_cmds))
    return t


def _bullets(items: list[str], s: dict) -> list:
    return [_p(f"• {item}", s["bullet"]) for item in items]


def _header_footer(canvas, doc, kicker: str) -> None:
    canvas.saveState()
    canvas.setFillColor(BRAND_DARK)
    canvas.rect(0, A4[1] - 1.15 * cm, A4[0], 1.15 * cm, fill=1, stroke=0)
    canvas.setFillColor(WHITE)
    canvas.setFont("Times-Bold", 8)
    canvas.drawString(2 * cm, A4[1] - 0.75 * cm, "SGE MUNICIPAL")
    canvas.setFont("Times-Roman", 8)
    canvas.drawRightString(A4[0] - 2 * cm, A4[1] - 0.75 * cm, kicker)
    canvas.setFillColor(LINE)
    canvas.rect(0, 0, A4[0], 1.1 * cm, fill=1, stroke=0)
    canvas.setFillColor(MUTED)
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(2 * cm, 0.45 * cm, f"{kicker}  ·  {DATA}")
    canvas.drawRightString(A4[0] - 2 * cm, 0.45 * cm, f"{doc.page}")
    canvas.restoreState()


def _cover_block(s: dict, title: str, subtitle: str, meta_rows: list[tuple[str, str]]) -> list:
    flow: list = []
    banner = Table(
        [
            [
                _p("SISTEMA DE GESTÃO ESCOLAR MUNICIPAL", s["cover_kicker"]),
            ],
            [_p(title, s["cover_title"])],
            [_p(subtitle, s["cover_sub"])],
        ],
        colWidths=[17 * cm],
    )
    banner.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), BRAND_DARK),
                ("LEFTPADDING", (0, 0), (-1, -1), 16),
                ("RIGHTPADDING", (0, 0), (-1, -1), 16),
                ("TOPPADDING", (0, 0), (0, 0), 18),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 18),
                ("TOPPADDING", (0, 1), (-1, -2), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    flow.append(banner)
    flow.append(Spacer(1, 14))
    meta_data = [
        [_p(f"<b>{k}</b>", s["meta"]), _p(v, s["meta"])] for k, v in meta_rows
    ]
    meta = Table(meta_data, colWidths=[4.5 * cm, 12.5 * cm])
    meta.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), SURFACE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LINEBELOW", (0, 0), (-1, -2), 0.3, LINE),
            ]
        )
    )
    flow.append(meta)
    flow.append(Spacer(1, 12))
    return flow


def _build_pdf(path: Path, kicker: str, story: list) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.6 * cm,
        title=kicker,
        author="SGE Municipal",
    )
    doc.build(
        story,
        onFirstPage=lambda c, d: _header_footer(c, d, kicker),
        onLaterPages=lambda c, d: _header_footer(c, d, kicker),
    )


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def _set_run_font(run, *, size=11, bold=False, italic=False, color=DOCX_INK, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def _docx_base() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("SGE Municipal  ·  documentos comerciais  ·  agosto/2026")
        _set_run_font(run, size=8, color=DOCX_MUTED)
    return doc


def _dx_p(doc: Document, text: str, *, size=11, bold=False, italic=False, color=DOCX_INK, align="justify", space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _dx_h(doc: Document, text: str, level=1):
    if level == 1:
        _dx_p(doc, text, size=14, bold=True, color=DOCX_BRAND, align="left", space_before=14, space_after=8)
    else:
        _dx_p(doc, text, size=12, bold=True, color=DOCX_BRAND, align="left", space_before=10, space_after=6)


def _dx_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        _set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def _dx_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = hdr[i]._tePr if False else hdr[i]._tc.get_or_add_tcPr()
        from docx.oxml import parse_xml

        shading.append(
            parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="152C4A"/>')
        )
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size=9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

ONDAS = [
    ["1", "P1-HARDEN", "Implementado", "n/a", "Compose de produção, recusa de boot inseguro, health/ready, CI deploy-check."],
    ["2", "P1-BACKUP", "Implementado", "n/a", "pg_dump noturno, retenção de 30 dias, disparo manual restrito à SME."],
    ["3", "P1-TRANSFER", "Implementado", "Implementado", "Encerra matrícula de origem, cria destino, respeita capacidade."],
    ["3", "P1-AUDIT", "Implementado", "Implementado", "AuditLog no banco; dashboard com atividade recente real."],
    ["4", "P1-LGPD", "Implementado", "Parcial", "Consentimento, portabilidade e anonimização. Falta aceite no cadastro."],
    ["4", "P2-PWDRESET", "Implementado", "Implementado", "Esqueci senha, redefinição por token e alteração na ficha de perfil."],
    ["4", "P2-DOC-UPLOAD", "Implementado", "Implementado", "Upload multipart com validação de MIME e isolamento por escola."],
    ["5", "P2-CLASS-CRUD", "Implementado", "Implementado", "Criação e edição de turmas e salas na interface."],
    ["5", "P2-USERS-UI", "Implementado", "Implementado", "SME cadastra diretores e secretários sem Django Admin."],
    ["5", "P2-REPORTS-UI", "Implementado", "Parcial", "Boletim e carteirinha em PDF. Falta histórico escolar em PDF."],
    ["6", "P2-EDUCACENSO", "Implementado", "Implementado", "Validação + ZIP estruturado. Não é o layout oficial do INEP."],
    ["6", "P2-GUARDIAN", "Implementado", "Implementado", "Portal Meus Filhos. CRUD de responsáveis ainda é placeholder."],
    ["6", "P2-NOTIF", "Implementado", "Implementado", "Notificações in-app com gatilhos. Sem e-mail/WhatsApp."],
    ["6", "P2-YEAREND", "Implementado", "Implementado", "Encerramento de ano, consolidação de histórico e trava do diário."],
]

COMPARATIVO = [
    ["Prontidão vs produção mínima SME", "~55%", "~88%"],
    ["Linhas de aplicação (backend + frontend)", "~23–25 mil", "~36 mil"],
    ["Casos de teste pytest", "~189", "~288"],
    ["Páginas React roteáveis", "núcleo + placeholders", "38 páginas (2 placeholders)"],
    ["Âncora Formato 1 (cessão as-is)", "R$ 320.000", "R$ 520.000"],
    ["Formato 2 (código + 12 meses)", "R$ 420.000", "R$ 540.000"],
    ["Piso interno de walk-away", "R$ 190.000", "R$ 340.000"],
    ["Custo estimado de reposição", "R$ 580.000", "R$ 780.000"],
]

EVOLUCOES = [
    ["SMTP de produção + templates de e-mail", "S", "R$ 5.500", "Necessário para reset de senha em produção."],
    ["Drill de restore e runbook de DR", "S", "R$ 5.500", "O backup existe; a restauração ainda precisa ser ensaiada."],
    ["Histórico escolar PDF + aceite LGPD no cadastro", "M", "R$ 14.000", "Gaps residuais do plano mínimo."],
    ["CRUD de responsáveis e conteúdo ministrado", "M", "R$ 14.000", "Telas hoje são placeholder."],
    ["WhatsApp (Twilio) para avisos à família", "M", "R$ 14.000", "Fora do mínimo; o canal in-app já existe."],
    ["Autenticação 2FA TOTP", "M", "R$ 14.000", "Fora do escopo do plano mínimo."],
    ["Educacenso em layout oficial INEP", "L", "R$ 65.000", "A exportação estruturada já está inclusa."],
    ["Assinatura digital ICP-Brasil", "L", "R$ 42.000", "PDFs atuais têm texto institucional e QR Code, sem certificado."],
    ["Merenda escolar / PNAE", "XL", "Sob consulta", "Fora de escopo."],
    ["Transporte escolar", "XL", "Sob consulta", "Fora de escopo."],
]

ESCOPO_INCLUSO = [
    "Cadastro único de alunos, escolas, turmas, salas, professores e alocação docente com prevenção de conflito de turno.",
    "Matrículas com regra de capacidade e bloqueio de matrícula ativa duplicada no ano letivo.",
    "Transferências entre unidades com efeito real: encerramento na origem e nova matrícula no destino.",
    "Diário de classe em lote (notas, frequência e pareceres descritivos da Educação Infantil).",
    "Fechamento de ano letivo com consolidação de histórico e trava de lançamentos.",
    "Boletim escolar em PDF, carteirinha com QR Code, exportações Excel/CSV e motor de consistência/exportação do Educacenso (formato estruturado próprio, não o leiaute oficial do INEP).",
    "Portal do responsável (visão Meus Filhos), mensagens internas e notificações in-app.",
    "Gestão de usuários da rede pela SME, recuperação de senha e alteração de perfil.",
    "Upload de documentos com isolamento por papel/escola e módulo mínimo de LGPD (consentimento, portabilidade e anonimização).",
    "Trilha de auditoria persistida, backup automatizado do banco e empacotamento de produção (Docker, Nginx, checagens de segurança).",
    "Carga inicial a partir do Censo Escolar 2025 do INEP (recorte municipal).",
    "Controle de acesso hierárquico (SME, direção, secretaria, docente, responsável).",
]

FORA_ESCOPO = [
    "Módulo de merenda escolar / PNAE.",
    "Módulo de transporte escolar e gestão de frotas.",
    "Módulo financeiro, empenhos ou folha de pagamento de servidores.",
    "Aplicativo nativo iOS/Android (o sistema é web responsivo).",
    "Autenticação de dois fatores (TOTP/hardware).",
    "Homologação formal com selo INEP/MEC e leiaute oficial do Educacenso.",
    "Assinatura digital com certificado ICP-Brasil.",
    "Integrações com sistemas de folha, protocolo ou WhatsApp Business.",
    "Hospedagem, domínio, certificados TLS e operação de infraestrutura do contratante, salvo o apoio ao primeiro deploy previsto no Formato 2.",
]


# ---------------------------------------------------------------------------
# Document 1 — Diagnóstico
# ---------------------------------------------------------------------------


def story_diagnostico(s: dict) -> list:
    story = _cover_block(
        s,
        "Diagnóstico de atendimento aos requisitos mínimos de produção",
        "Reavaliação independente do repositório contra o plano executivo da SME.",
        [
            ("Documento", "Diagnóstico técnico-comercial"),
            ("Referência", REF.replace("PC", "DX")),
            ("Data", DATA),
            ("Objeto", "Sistema de Gestão Escolar Municipal (SGE)"),
            ("Plano de referência", "PLANO_EXECUCAO_PRODUCAO_MINIMA v1.0.0"),
            ("Método", "Auditoria no código (services, APIs, UI e testes), não em STATUS.md"),
        ],
    )
    story.append(
        _p(
            "Veredito: o repositório ATENDE os requisitos mínimos de produção definidos para Secretarias Municipais de Educação. As seis ondas do plano (P1 e P2) estão implementadas no backend. No frontend, dez dos doze itens de interface estão completos e dois permanecem parciais (LGPD no cadastro e emissão de histórico escolar em PDF). Nenhum item do plano está ausente.",
            s["verdict"],
        )
    )
    story.append(
        _p(
            "Esta conclusão inverte a avaliação comercial anterior, de agosto de 2026, que classificava o produto como MVP avançado do núcleo pedagógico (~55% de prontidão), com auditoria, backup, Educacenso, portal da família e notificações ainda em esqueleto. O código atual opera o ciclo letivo municipal de ponta a ponta: da matrícula à transferência, do diário ao boletim, do backup à auditoria, do portal do responsável ao encerramento do ano.",
            s["body"],
        )
    )
    story.append(_p("1. Critério utilizado", s["h1"]))
    story.append(
        _p(
            "Foi utilizado exclusivamente o Definition of Done e as fichas técnicas do plano de produção mínima. Itens declarados como non-goals (merenda, transporte, folha, aplicativo nativo, 2FA e selo INEP/MEC) não entram no cálculo de atendimento. Documentos internos desatualizados (STATUS.md de julho/2026 e trechos do README que ainda mencionam auditoria só em log) foram ignorados quando contradiziam o código.",
            s["body"],
        )
    )
    story.append(_p("2. Evolução desde a avaliação anterior", s["h1"]))
    story.append(
        _table(
            ["Indicador", "Avaliação anterior", "30/08/2026"],
            COMPARATIVO,
            [6.2 * cm, 5.4 * cm, 5.4 * cm],
            s,
        )
    )
    story.append(
        _p(
            "Fonte métrica: 21.015 linhas Python (excluídas migrations), 15.295 linhas TypeScript em frontend/src, 39 arquivos de teste de backend, cerca de 288 casos pytest, 38 páginas React e 2 rotas ainda placeholder (Responsáveis e Conteúdo ministrado).",
            s["small"],
        )
    )
    story.append(_p("3. Atendimento por onda", s["h1"]))
    story.append(
        _table(
            ["Onda", "Item", "Backend", "Frontend", "Observação"],
            ONDAS,
            [1.4 * cm, 3.2 * cm, 2.6 * cm, 2.6 * cm, 7.2 * cm],
            s,
        )
    )
    story.append(_p("4. Checklist de go-live do plano", s["h1"]))
    story.append(
        _table(
            ["Item de go-live", "Situação", "Comentário"],
            [
                ["Segurança (DEBUG, SECRET_KEY, cookies, portas internas)", "Atende", "Recusa de boot inseguro em ENVIRONMENT=production. Compose de produção não publica 5432/6379."],
                ["Backups automatizados", "Atende, com ressalva", "Dump noturno e retenção existem. Falta drill de restauração em homologação."],
                ["Transferência escolar ponta a ponta", "Atende", "Origem encerrada, destino criado, capacidade com lock pessimista."],
                ["Auditoria persistida + LGPD mínimo", "Atende, com ressalva", "Trilha no banco. Aceite de termos ainda não é obrigatório no cadastro."],
                ["Operação da SME sem Django Admin", "Atende", "Turmas, salas, professores e usuários da rede pela interface."],
                ["Documentos oficiais (boletim e carteirinha)", "Atende, com ressalva", "PDFs oficiais na UI. Histórico escolar em PDF ainda não tem botão na ficha."],
                ["Consistência do Censo / Educacenso", "Atende o mínimo do plano", "Validação + ZIP estruturado. Não é o leiaute oficial INEP (non-goal)."],
                ["Portal da família", "Atende", "Responsável com múltiplos filhos; boletim por dependente."],
                ["Encerramento de ano letivo", "Atende", "Consolidação de histórico e trava do diário."],
            ],
            [5.5 * cm, 3.3 * cm, 8.2 * cm],
            s,
        )
    )
    story.append(_p("5. Gaps residuais (não bloqueantes para piloto)", s["h1"]))
    story.extend(
        _bullets(
            [
                "Aceite de termos LGPD obrigatório no cadastro/matrícula (hoje o consentimento vive na ficha do aluno).",
                "Emissão de histórico escolar em PDF a partir da ficha do aluno.",
                "Telas de Responsáveis e Conteúdo ministrado ainda são placeholder.",
                "Ensaio de restore do backup em ambiente de homologação.",
                "SMTP de produção: em desenvolvimento o e-mail de reset cai no console.",
                "CI ainda marca lint, formatação, mypy, Bandit e Sonar como continue-on-error (os testes pytest não).",
                "Licença MIT declarada no pyproject.toml — incompatível com cessão exclusiva enquanto o repositório for público.",
            ],
            s,
        )
    )
    story.append(Spacer(1, 6))
    story.append(_p("6. O que continua fora de escopo", s["h1"]))
    story.extend(_bullets(FORA_ESCOPO, s))
    story.append(_p("7. Recomendação", s["h1"]))
    story.append(
        _p(
            "O produto está apto a um piloto municipal homologado internamente (secretaria + recorte de escolas), com capacitação e apoio de implantação. Não é um ERP educacional completo nem um sistema homologado pelo INEP. A recomendação comercial é ofertar o Formato 2 (cessão + 12 meses) à prefeitura, e o Formato 1 (cessão fechada) apenas a comprador com time próprio de engenharia. Os gaps residuais cabem em itens S/M do catálogo de evoluções e não justificam mais o desconto de MVP aplicado na avaliação anterior.",
            s["body"],
        )
    )
    story.append(
        _p(
            "Este diagnóstico não substitui teste de aceitação da secretaria nem auditoria de segurança contratada. É uma leitura de prontidão funcional com base no código existente nesta data.",
            s["small"],
        )
    )
    return story


def docx_diagnostico() -> Document:
    doc = _docx_base()
    _dx_p(doc, "SISTEMA DE GESTÃO ESCOLAR MUNICIPAL", size=10, bold=True, color=DOCX_BRAND, align="left")
    _dx_p(doc, "Diagnóstico de atendimento aos requisitos mínimos de produção", size=18, bold=True, color=DOCX_INK, align="left", space_after=4)
    _dx_p(doc, "Reavaliação independente do repositório contra o plano executivo da SME.", size=11, italic=True, color=DOCX_MUTED, align="left")
    _dx_p(doc, f"Referência DX-SGE-001/2026  ·  {DATA}  ·  Plano: PLANO_EXECUCAO_PRODUCAO_MINIMA v1.0.0", size=9, color=DOCX_MUTED, align="left")
    _dx_p(
        doc,
        "Veredito: o repositório ATENDE os requisitos mínimos de produção definidos para Secretarias Municipais de Educação. As seis ondas do plano (P1 e P2) estão implementadas no backend. No frontend, dez dos doze itens de interface estão completos e dois permanecem parciais. Nenhum item do plano está ausente.",
        bold=True,
        color=DOCX_OK,
        align="left",
    )
    _dx_p(
        doc,
        "Esta conclusão inverte a avaliação comercial anterior, que classificava o produto como MVP avançado do núcleo pedagógico (~55% de prontidão). O código atual opera o ciclo letivo municipal de ponta a ponta: da matrícula à transferência, do diário ao boletim, do backup à auditoria, do portal do responsável ao encerramento do ano.",
    )
    _dx_h(doc, "1. Critério utilizado")
    _dx_p(
        doc,
        "Foi utilizado exclusivamente o Definition of Done e as fichas técnicas do plano de produção mínima. Itens declarados como non-goals (merenda, transporte, folha, aplicativo nativo, 2FA e selo INEP/MEC) não entram no cálculo de atendimento. Documentos internos desatualizados foram ignorados quando contradiziam o código.",
    )
    _dx_h(doc, "2. Evolução desde a avaliação anterior")
    _dx_table(doc, ["Indicador", "Avaliação anterior", "30/08/2026"], COMPARATIVO)
    _dx_p(
        doc,
        "Fonte: 21.015 linhas Python (sem migrations), 15.295 linhas TypeScript, ~288 casos pytest, 38 páginas React, 2 placeholders.",
        size=9,
        italic=True,
        color=DOCX_MUTED,
    )
    _dx_h(doc, "3. Atendimento por onda")
    _dx_table(doc, ["Onda", "Item", "Backend", "Frontend", "Observação"], ONDAS)
    _dx_h(doc, "4. Gaps residuais (não bloqueantes para piloto)")
    _dx_bullets(
        doc,
        [
            "Aceite de termos LGPD obrigatório no cadastro/matrícula.",
            "Emissão de histórico escolar em PDF na ficha do aluno.",
            "Telas de Responsáveis e Conteúdo ministrado ainda são placeholder.",
            "Ensaio de restore do backup em homologação.",
            "SMTP de produção para o fluxo de reset de senha.",
            "CI com lint/mypy/Bandit/Sonar em continue-on-error.",
            "Licença MIT no pyproject.toml — incompatível com cessão exclusiva se o repositório for público.",
        ],
    )
    _dx_h(doc, "5. Recomendação")
    _dx_p(
        doc,
        "O produto está apto a um piloto municipal homologado internamente. Não é um ERP educacional completo nem um sistema homologado pelo INEP. Recomenda-se ofertar o Formato 2 à prefeitura e o Formato 1 apenas a comprador com engenharia própria. Os gaps residuais não justificam mais o desconto de MVP da avaliação anterior.",
    )
    return doc


# ---------------------------------------------------------------------------
# Document 2 — Precificação (interno)
# ---------------------------------------------------------------------------


def story_precificacao(s: dict) -> list:
    story = _cover_block(
        s,
        "Memorando interno de precificação",
        "Cessão de código-fonte e contrato de 12 meses — valores reprecificados após a produção mínima.",
        [
            ("Documento", "Memorando interno — não enviar ao comprador"),
            ("Referência", "PF-SGE-002/2026"),
            ("Data", DATA),
            ("Público", "Sócios / direção comercial"),
            ("Validade interna", "90 dias ou até nova reavaliação de código"),
        ],
    )
    story.append(
        _p(
            "O produto deixou de ser um MVP. A âncora anterior de R$ 320 mil pressupunha que o comprador herdaria cerca de R$ 105 a 200 mil em evoluções obrigatórias (transferência real, auditoria, LGPD, Educacenso estruturado, portal, fechamento de ano, UI operacional). Esse pacote foi absorvido pelo fonte. O preço sobe porque o risco caiu, não porque se inventou um módulo novo de merenda.",
            s["body"],
        )
    )
    story.append(_p("1. Números-âncora", s["h1"]))
    story.append(
        _table(
            ["Peça", "Valor", "Leitura"],
            [
                ["Formato 1 — lista (cessão exclusiva as-is)", "R$ 520.000", "Faixa R$ 420–620 mil. Piso de walk-away: R$ 340 mil."],
                ["Formato 2 — padrão (código + 12 meses)", "R$ 540.000", "R$ 360 mil de cessão + 12 × R$ 15 mil. À vista: R$ 513 mil."],
                ["Ano 2 do Formato 2", "R$ 16.500 / mês", "Sem nova cessão. Reajuste IPCA + 2 p.p."],
                ["Custo de reposição", "R$ 780.000", "12–14 meses, ~2,5 FTE, software house reconstruindo o produto atual."],
                ["Âncora de mercado (i-Educar SLM-PE 2021)", "R$ 347 mil / ano", "Inflacionado ~R$ 500 mil em 2026; produto público e mais maduro."],
            ],
            [6.5 * cm, 3.8 * cm, 6.7 * cm],
            s,
        )
    )
    story.append(_p("2. Os dois formatos", s["h1"]))
    story.append(_p("Formato 1 — cessão fechada", s["h2"]))
    story.append(
        _p(
            "Inclui código-fonte backend e frontend, Docker Compose de produção, documentação técnica, OpenAPI, 40 horas de transferência de conhecimento e 90 dias de correção de defeitos já existentes na entrega. Não inclui hospedagem, deploy na infra do cliente, treinamento da rede, SMTP, drill de restore nem feature nova. Pagamento: 40% na assinatura, 40% na entrega do repositório, 20% 30 dias após o handover.",
            s["body"],
        )
    )
    story.append(_p("Formato 2 — código + 12 meses (proposta padrão)", s["h2"]))
    story.append(
        _p(
            "Tudo do Formato 1, mais bolsa de 32 horas/mês (384 h/ano) para manutenção corretiva, patches de segurança e evoluções S/M; SLA 8×5 (P1 em 8 horas úteis, P2 em 24 horas); 16 horas de capacitação; apoio ao primeiro deploy. Hora extra da bolsa: R$ 220 (versus R$ 250 à la carte). Módulos L/XL com 12% de desconto sobre a tabela. Condições: R$ 360 mil na assinatura + mensalidade até o dia 10. Até 16 horas não usadas rolam para o mês seguinte.",
            s["body"],
        )
    )
    story.append(_p("3. Gasto esperado no ano 1", s["h1"]))
    story.append(
        _table(
            ["Cenário", "Formato 1", "Formato 2"],
            [
                ["Só o código (zip + handover)", "R$ 520 mil", "R$ 540 mil"],
                ["Piloto em produção (SMTP, restore, histórico PDF)", "R$ 545 mil", "R$ 540 mil (cabe na bolsa)"],
                ["Rede com extra L (Educacenso layout INEP)", "R$ 585 mil", "R$ 597 mil (R$ 65 mil com 12% off = ~R$ 57 mil)"],
            ],
            [8 * cm, 4.5 * cm, 4.5 * cm],
            s,
        )
    )
    story.append(
        _p(
            "No cenário piloto, o Formato 2 sai mais barato e previsível. No cenário “só quero o zip”, o Formato 1 ganha — e o comprador herda o risco residual.",
            s["small"],
        )
    )
    story.append(_p("4. Como fechar por tipo de comprador", s["h1"]))
    story.append(
        _p(
            "<b>Prefeitura / SME.</b> Formato 2. Sem time de engenharia, o zip não opera a rede. O produto agora aguenta piloto; o contrato carrega deploy, SMTP, capacitação e os gaps S/M.",
            s["body"],
        )
    )
    story.append(
        _p(
            "<b>Software house / ISV.</b> Formato 1 a R$ 520 mil (chão R$ 420 mil à vista). Eles têm engenharia para o residual. O vendedor vira parceiro de horas L/XL, não subcontratado em retainer cheio.",
            s["body"],
        )
    )
    story.append(
        _p(
            "<b>Investidor / parceiro.</b> Valuation de reposição ~R$ 780 mil; âncora de cessão R$ 520 mil. Diluição faz sentido se vier distribuição (vários municípios) ou módulo XL. Não vender abaixo de R$ 340 mil.",
            s["body"],
        )
    )
    story.append(_p("5. Catálogo de evoluções que ainda restam", s["h1"]))
    story.append(
        _p(
            "Itens da avaliação anterior que saíram do catálogo porque já estão no produto: transferência com efeito real, auditoria persistida, LGPD mínimo, Educacenso estruturado, portal do responsável, notificações in-app, fechamento de ano, CRUD de turmas/usuários, upload de documentos e recuperação de senha.",
            s["body"],
        )
    )
    story.append(
        _table(
            ["Item", "Porte", "Lista", "Nota"],
            EVOLUCOES,
            [5.8 * cm, 1.6 * cm, 2.4 * cm, 7.2 * cm],
            s,
        )
    )
    story.append(
        _p(
            "Tabela de portes: XS 4–8 h R$ 1.500 · S 16–24 h R$ 5.500 · M 40–80 h R$ 14.000 · L 120–200 h R$ 42.000 · XL 320 h+ sob consulta. Urgente +28% (R$ 320/h).",
            s["small"],
        )
    )
    story.append(_p("6. Condição suspensiva de IP", s["h1"]))
    story.append(
        _p(
            "O pyproject.toml ainda declara licença MIT e o README aponta GitHub. Se o repositório for público, não existe cessão exclusiva. Antes de assinar: tornar o remoto privado, retirar a MIT, ceder direitos patrimoniais no contrato. Se o código já vazou como MIT, o que se vende é know-how + contrato de implantação — não o monopólio do fonte.",
            s["warn"],
        )
    )
    story.append(
        _p(
            "Uso: este memorando é interno. Ao comprador, envie a Proposta Comercial oficial (PC-SGE-001/2026), não este arquivo.",
            s["small"],
        )
    )
    return story


def docx_precificacao() -> Document:
    doc = _docx_base()
    _dx_p(doc, "MEMORANDO INTERNO — NÃO ENVIAR AO COMPRADOR", size=9, bold=True, color=DOCX_WARN, align="left")
    _dx_p(doc, "Precificação comercial — SGE Municipal", size=18, bold=True, color=DOCX_INK, align="left", space_after=4)
    _dx_p(doc, f"PF-SGE-002/2026  ·  {DATA}  ·  uso exclusivo de sócios e direção comercial", size=9, color=DOCX_MUTED, align="left")
    _dx_p(
        doc,
        "O produto deixou de ser um MVP. A âncora anterior de R$ 320 mil pressupunha evoluções obrigatórias que agora estão no código. O preço sobe porque o risco caiu.",
    )
    _dx_h(doc, "1. Números-âncora")
    _dx_table(
        doc,
        ["Peça", "Valor", "Leitura"],
        [
            ["Formato 1 — lista", "R$ 520.000", "Faixa R$ 420–620 mil. Piso: R$ 340 mil."],
            ["Formato 2 — padrão", "R$ 540.000", "R$ 360 mil + 12 × R$ 15 mil. À vista: R$ 513 mil."],
            ["Ano 2 Formato 2", "R$ 16.500/mês", "Sem nova cessão. IPCA + 2 p.p."],
            ["Custo de reposição", "R$ 780.000", "12–14 meses, ~2,5 FTE."],
            ["Âncora i-Educar SLM-PE 2021", "R$ 347 mil/ano", "~R$ 500 mil em 2026."],
        ],
    )
    _dx_h(doc, "2. Formato 1 — cessão fechada")
    _dx_p(
        doc,
        "Código, Docker de produção, docs, OpenAPI, 40 h de handover, 90 dias de defeito já existente. Sem hospedagem, treino nem feature nova. 40/40/20.",
    )
    _dx_h(doc, "3. Formato 2 — código + 12 meses")
    _dx_p(
        doc,
        "Tudo do Formato 1 + 32 h/mês, SLA 8×5, 16 h de capacitação, primeiro deploy. Hora extra R$ 220. L/XL com 12% off.",
    )
    _dx_h(doc, "4. Como fechar")
    _dx_bullets(
        doc,
        [
            "Prefeitura / SME: Formato 2 (R$ 540 mil no ano 1).",
            "Software house: Formato 1 a R$ 520 mil (chão R$ 420 mil à vista).",
            "Investidor: não vender abaixo de R$ 340 mil.",
        ],
    )
    _dx_h(doc, "5. Catálogo residual")
    _dx_table(doc, ["Item", "Porte", "Lista", "Nota"], EVOLUCOES)
    _dx_h(doc, "6. IP")
    _dx_p(
        doc,
        "MIT no pyproject.toml impede cessão exclusiva enquanto o repositório for público. Fechar o remoto e alterar a licença antes da assinatura.",
        color=DOCX_WARN,
    )
    return doc


# ---------------------------------------------------------------------------
# Document 3 — Proposta comercial oficial
# ---------------------------------------------------------------------------


def story_proposta(s: dict) -> list:
    story = _cover_block(
        s,
        "Proposta comercial oficial",
        "Cessão de direitos patrimoniais sobre o Sistema de Gestão Escolar Municipal e, opcionalmente, serviços de implantação e suporte pelo prazo de 12 meses.",
        [
            ("Proposta nº", REF),
            ("Data de emissão", DATA),
            ("Validade", VALIDADE),
            ("Proponente", "[Razão social]  ·  CNPJ [00.000.000/0001-00]"),
            ("Destinatário", "Secretaria Municipal de Educação de [Município/UF]"),
            ("Objeto", "SGE Municipal — produção mínima operacional"),
            ("Moeda", "Real (R$)"),
        ],
    )
    story.append(_p("1. Objeto", s["h1"]))
    story.append(
        _p(
            "A presente proposta tem por objeto a cessão onerosa dos direitos patrimoniais de uso e exploração do software Sistema de Gestão Escolar Municipal (SGE), incluindo códigos-fonte do backend e do frontend, artefatos de containerização, documentação técnica e, conforme a opção contratada, serviços de implantação, capacitação e suporte pelo prazo de 12 (doze) meses.",
            s["body"],
        )
    )
    story.append(
        _p(
            "O SGE é uma plataforma web destinada a Secretarias Municipais de Educação para gestão da rede pública de ensino: cadastro único de alunos, escolas, turmas, diário de classe, matrículas e transferências, matriz curricular, documentos oficiais, portal do responsável, auditoria, backup e encerramento de ano letivo. O controle de acesso é hierárquico por papel e por unidade escolar.",
            s["body"],
        )
    )
    story.append(_p("2. Fundamentação e estado do produto", s["h1"]))
    story.append(
        _p(
            "O software encontra-se em estado de produção mínima operacional, nos termos do plano técnico de go-live municipal. Em números: arquitetura Django 6.1 + React 18, aproximadamente 36 mil linhas de aplicação, cerca de 288 casos de teste automatizado de backend, empacotamento Docker de produção (PostgreSQL 16, Redis, Gunicorn, Nginx) e carga inicial a partir do Censo Escolar 2025 do INEP. O núcleo pedagógico e a operação da secretaria (usuários, turmas, relatórios, Educacenso estruturado, portal da família e fechamento de ano) estão implementados e utilizáveis em piloto.",
            s["body"],
        )
    )
    story.append(
        _p(
            "A solução não se confunde com sistemas de merenda, transporte, folha de pagamento ou com o Educacenso homologado pelo INEP. Esses módulos, quando de interesse do contratante, são cotados em separado.",
            s["body"],
        )
    )
    story.append(_p("3. Escopo incluso", s["h1"]))
    story.extend(_bullets(ESCOPO_INCLUSO, s))
    story.append(_p("4. Fora de escopo", s["h1"]))
    story.extend(_bullets(FORA_ESCOPO, s))
    story.append(_p("5. Opções comerciais", s["h1"]))
    story.append(
        _p(
            "O contratante escolhe uma das opções abaixo. Não são cumulativas. A Opção B é a recomendada para órgãos da Administração Pública Municipal sem equipe própria de engenharia de software.",
            s["body"],
        )
    )
    story.append(_p("5.1. Opção A — Cessão do código-fonte", s["h2"]))
    story.append(
        _table(
            ["Rubrica", "Condição"],
            [
                ["Preço global", "R$ 520.000,00 (quinhentos e vinte mil reais)"],
                ["Faixa de negociação", "R$ 420.000,00 a R$ 620.000,00, mediante justificativa de escopo"],
                ["Incluso", "Códigos-fonte, Docker de produção, documentação técnica, OpenAPI, 40 horas de transferência de conhecimento, 90 dias de garantia de defeitos existentes na entrega"],
                ["Não incluso", "Hospedagem, deploy, capacitação da rede, novas funcionalidades, SMTP, operação contínua"],
                ["Pagamento", "40% na assinatura · 40% na entrega do repositório e da documentação · 20% 30 dias após o handover"],
            ],
            [4.2 * cm, 12.8 * cm],
            s,
        )
    )
    story.append(_p("5.2. Opção B — Cessão + 12 meses de serviços (padrão)", s["h2"]))
    story.append(
        _table(
            ["Rubrica", "Condição"],
            [
                ["Preço do ano 1", "R$ 540.000,00 (quinhentos e quarenta mil reais)"],
                ["Composição", "R$ 360.000,00 pela cessão + 12 parcelas de R$ 15.000,00"],
                ["À vista", "R$ 513.000,00 (desconto de 5% sobre o global do ano 1)"],
                ["Incluso além da Opção A", "32 horas/mês (384 h/ano) de manutenção, patches e evoluções de pequeno/médio porte; SLA 8×5; 16 horas de capacitação; apoio ao primeiro deploy na infraestrutura do contratante"],
                ["Hora excedente", "R$ 220,00"],
                ["Módulos grandes (L/XL)", "Cotação à parte, com 12% de desconto sobre a tabela de evoluções"],
                ["Ano 2 em diante", "R$ 16.500,00 / mês, sem nova cessão, reajuste IPCA + 2 p.p. no mês 13"],
                ["Pagamento", "R$ 360.000,00 na assinatura · mensalidades até o dia 10 de cada mês · até 16 horas não usadas transferem-se ao mês seguinte"],
            ],
            [4.2 * cm, 12.8 * cm],
            s,
        )
    )
    story.append(_p("5.3. Tabela de evoluções extraordinárias", s["h2"]))
    story.append(
        _p(
            "Demandas fora da bolsa da Opção B, ou qualquer evolução na Opção A, seguem portes predefinidos, para fins de planejamento orçamentário e, quando aplicável, de aditivo:",
            s["body"],
        )
    )
    story.append(
        _table(
            ["Porte", "Esforço estimado", "Preço"],
            [
                ["XS", "4 a 8 horas", "R$ 1.500,00"],
                ["S", "16 a 24 horas", "R$ 5.500,00"],
                ["M", "40 a 80 horas", "R$ 14.000,00"],
                ["L", "120 a 200 horas", "R$ 42.000,00"],
                ["XL", "320 horas ou mais", "Sob consulta"],
                ["Urgência", "Fora da janela 8×5", "Acréscimo de 28% (equivalente a R$ 320,00/hora)"],
            ],
            [4 * cm, 6.5 * cm, 6.5 * cm],
            s,
        )
    )
    story.append(Spacer(1, 6))
    story.append(
        _table(
            ["Evolução típica residual", "Porte", "Preço de lista"],
            [[r[0], r[1], r[2]] for r in EVOLUCOES],
            [10.2 * cm, 2.4 * cm, 4.4 * cm],
            s,
        )
    )
    story.append(_p("6. Entregáveis", s["h1"]))
    story.extend(
        _bullets(
            [
                "Repositório privado contendo backend, frontend, Docker Compose de produção, exemplos de variáveis de ambiente e documentação de arquitetura.",
                "Acesso à especificação OpenAPI da API REST.",
                "Pacote de carga inicial a partir do Censo Escolar (quando o contratante fornecer ou autorizar o recorte municipal).",
                "Na Opção B: plano de implantação de 12 semanas, atas de capacitação e canal de suporte no horário 8×5.",
            ],
            s,
        )
    )
    story.append(_p("7. Cronograma de implantação (Opção B)", s["h1"]))
    story.append(
        _table(
            ["Semanas", "Atividade", "Responsável"],
            [
                ["1 e 2", "Provisionamento (TLS, DNS, SMTP, secrets), migrate, backup e ensaio de restore, carga do Censo ou base piloto", "Contratada + TI municipal"],
                ["3 e 4", "Parametrização: ano letivo, etapas, matrizes, usuários da SME e das escolas piloto", "Contratada + SME"],
                ["5 a 8", "Capacitação por papel (SME, direção, secretaria, docente, responsável) e operação assistida em 2 a 3 unidades", "Contratada + escolas piloto"],
                ["9 a 12", "Go-live do recorte acordado, acompanhamento, correções e relatório de encerramento da implantação", "Contratada + SME"],
            ],
            [3 * cm, 9.5 * cm, 4.5 * cm],
            s,
        )
    )
    story.append(
        _p(
            "Prazos contam-se a partir da assinatura, do pagamento da parcela inicial e da disponibilização, pelo contratante, de infraestrutura mínima (servidor ou nuvem, domínio, caixa postal SMTP e interlocutor da SME). A Opção A não inclui este cronograma, salvo as 40 horas de handover.",
            s["small"],
        )
    )
    story.append(_p("8. Níveis de serviço (Opção B)", s["h1"]))
    story.append(
        _table(
            ["Severidade", "Definição", "Prazo de resposta", "Prazo de solução paliativa"],
            [
                ["P1", "Indisponibilidade total do sistema em produção", "4 horas úteis", "8 horas úteis"],
                ["P2", "Função crítica degradada (matrícula, diário, login)", "8 horas úteis", "24 horas úteis"],
                ["P3", "Falha não bloqueante ou dúvida operacional", "16 horas úteis", "5 dias úteis"],
                ["P4", "Evolução ou melhoria", "2 dias úteis", "Conforme porte da tabela"],
            ],
            [2.4 * cm, 6.2 * cm, 4.2 * cm, 4.2 * cm],
            s,
        )
    )
    story.append(
        _p(
            "Horário de cobertura: segunda a sexta, das 8h às 18h (horário de Brasília), excluídos feriados nacionais. Fora dessa janela aplica-se o acréscimo de urgência, se demandado pelo contratante.",
            s["body"],
        )
    )
    story.append(_p("9. Propriedade intelectual e confidencialidade", s["h1"]))
    story.append(
        _p(
            "Com o pagamento integral da cessão, o contratante recebe os direitos patrimoniais de uso, modificação e exploração do SGE no âmbito de sua rede municipal, incluindo o direito de contratar terceiros para manutenção. O proponente retém o direito moral de autoria e, salvo cláusula de sigilo em contrário, o direito de citar o município como caso de uso.",
            s["body"],
        )
    )
    story.append(
        _p(
            "A cessão pressupõe repositório privado e licença proprietária no ato da entrega. O contratante não poderá republicar o código sob licença livre sem autorização escrita. Dados pessoais de alunos, responsáveis e servidores permanecem de titularidade do município, nos termos da Lei nº 13.709/2018 (LGPD).",
            s["body"],
        )
    )
    story.append(_p("10. Garantia", s["h1"]))
    story.append(
        _p(
            "A contratada garante, pelo prazo de 90 (noventa) dias contados da entrega do repositório, a correção de defeitos comprovadamente existentes no software na data da entrega, desde que reproduzíveis em ambiente controlado. A garantia não cobre customizações feitas por terceiros, má operação, falhas de infraestrutura do contratante nem requisitos novos.",
            s["body"],
        )
    )
    story.append(_p("11. Obrigações do contratante", s["h1"]))
    story.extend(
        _bullets(
            [
                "Indicar gestor do contrato e interlocutores da SME e da TI municipal.",
                "Disponibilizar infraestrutura, certificados digitais de sítio (TLS), DNS e serviço de correio eletrônico (SMTP) para recuperação de senha e avisos.",
                "Obter as bases autorizadas do Censo ou equivalentes para a carga inicial.",
                "Assegurar o tratamento de dados pessoais na qualidade de controlador, nos termos da LGPD.",
                "Na Opção B, liberar os usuários-piloto para as sessões de capacitação nas datas acordadas.",
            ],
            s,
        )
    )
    story.append(_p("12. Condições gerais", s["h1"]))
    story.extend(
        _bullets(
            [
                f"Validade desta proposta: {VALIDADE}, contados da data de emissão.",
                "Preços em reais, sem incidência de ISS, PIS, Cofins ou IR já destacados; tributos eventualmente devidos no município do tomador serão acrescidos na nota fiscal conforme a legislação aplicável.",
                "A contratação com a Administração Pública observará a Lei nº 14.133/2021, quando couber, e poderá ser formalizada por inexigibilidade, credenciamento, dispensa ou licitação, conforme enquadramento jurídico do órgão.",
                "Foro: comarca do município contratante, salvo disposição em contrário no contrato.",
                "Esta proposta não gera obrigação de contratar até a assinatura do instrumento jurídico correspondente.",
            ],
            s,
        )
    )
    story.append(_p("13. Aceite", s["h1"]))
    story.append(
        _p(
            "O destinatário manifesta interesse na opção abaixo e devolve uma via assinada (física ou com assinatura eletrônica qualificada/avançada) para elaboração do contrato:",
            s["body"],
        )
    )
    story.append(
        _table(
            ["Opção", "Descrição", "Valor do ano 1", "Assinale"],
            [
                ["A", "Cessão do código-fonte", "R$ 520.000,00", ""],
                ["B", "Cessão + 12 meses de serviços", "R$ 540.000,00", ""],
                ["B à vista", "Opção B com 5% de desconto", "R$ 513.000,00", ""],
            ],
            [2.6 * cm, 6.4 * cm, 4.4 * cm, 3.6 * cm],
            s,
        )
    )
    story.append(Spacer(1, 18))
    story.append(
        _p(
            "Local e data: ________________________________, ____ de _________________ de 20____.",
            s["center"],
        )
    )
    story.append(Spacer(1, 22))
    sign = Table(
        [
            [
                _p("________________________________<br/><b>Proponente</b><br/>Nome:<br/>CPF/CNPJ:", s["center"]),
                _p("________________________________<br/><b>Destinatário</b><br/>Nome:<br/>Cargo / matrícula:", s["center"]),
            ]
        ],
        colWidths=[8.5 * cm, 8.5 * cm],
    )
    sign.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(sign)
    story.append(Spacer(1, 16))
    story.append(
        _p(
            "Anexos sugeridos ao contrato: (I) este instrumento de proposta; (II) diagnóstico de prontidão; (III) documentação de arquitetura; (IV) lista de papéis e permissões; (V) termo de tratamento de dados pessoais.",
            s["small"],
        )
    )
    return story


def docx_proposta() -> Document:
    doc = _docx_base()
    _dx_p(doc, "PROPOSTA COMERCIAL OFICIAL", size=10, bold=True, color=DOCX_BRAND, align="left")
    _dx_p(doc, "Sistema de Gestão Escolar Municipal (SGE)", size=18, bold=True, color=DOCX_INK, align="left", space_after=4)
    _dx_p(
        doc,
        f"Proposta nº {REF}  ·  Emissão: {DATA}  ·  Validade: {VALIDADE}",
        size=9,
        color=DOCX_MUTED,
        align="left",
    )
    _dx_p(doc, "Proponente: [Razão social]  ·  CNPJ [00.000.000/0001-00]", size=11, align="left")
    _dx_p(doc, "Destinatário: Secretaria Municipal de Educação de [Município/UF]", size=11, align="left")
    _dx_h(doc, "1. Objeto")
    _dx_p(
        doc,
        "Cessão onerosa dos direitos patrimoniais de uso e exploração do software SGE, incluindo códigos-fonte, artefatos de containerização, documentação técnica e, conforme a opção, serviços de implantação, capacitação e suporte por 12 meses.",
    )
    _dx_h(doc, "2. Estado do produto")
    _dx_p(
        doc,
        "O software encontra-se em estado de produção mínima operacional. Stack Django 6.1 + React 18, ~36 mil linhas de aplicação, ~288 testes de backend, Docker de produção e carga a partir do Censo Escolar 2025. O núcleo pedagógico e a operação da secretaria estão implementados e utilizáveis em piloto. Merenda, transporte, folha e Educacenso homologado pelo INEP permanecem fora de escopo.",
    )
    _dx_h(doc, "3. Escopo incluso")
    _dx_bullets(doc, ESCOPO_INCLUSO)
    _dx_h(doc, "4. Fora de escopo")
    _dx_bullets(doc, FORA_ESCOPO)
    _dx_h(doc, "5. Opções comerciais")
    _dx_p(
        doc,
        "A Opção B é a recomendada para a Administração Pública Municipal sem equipe própria de engenharia.",
    )
    _dx_h(doc, "5.1. Opção A — Cessão do código-fonte", 2)
    _dx_table(
        doc,
        ["Rubrica", "Condição"],
        [
            ["Preço global", "R$ 520.000,00"],
            ["Incluso", "Fonte, Docker de produção, docs, OpenAPI, 40 h de handover, 90 dias de garantia de defeitos existentes"],
            ["Não incluso", "Hospedagem, deploy, capacitação, novas funcionalidades, operação contínua"],
            ["Pagamento", "40% assinatura · 40% entrega · 20% 30 dias após o handover"],
        ],
    )
    _dx_h(doc, "5.2. Opção B — Cessão + 12 meses (padrão)", 2)
    _dx_table(
        doc,
        ["Rubrica", "Condição"],
        [
            ["Preço do ano 1", "R$ 540.000,00 (R$ 360.000,00 + 12 × R$ 15.000,00)"],
            ["À vista", "R$ 513.000,00 (5% de desconto)"],
            ["Incluso além da A", "32 h/mês, SLA 8×5, 16 h de capacitação, primeiro deploy"],
            ["Hora excedente", "R$ 220,00"],
            ["Ano 2", "R$ 16.500,00 / mês, IPCA + 2 p.p."],
        ],
    )
    _dx_h(doc, "5.3. Evoluções extraordinárias", 2)
    _dx_table(
        doc,
        ["Porte", "Esforço", "Preço"],
        [
            ["XS", "4 a 8 h", "R$ 1.500,00"],
            ["S", "16 a 24 h", "R$ 5.500,00"],
            ["M", "40 a 80 h", "R$ 14.000,00"],
            ["L", "120 a 200 h", "R$ 42.000,00"],
            ["XL", "320 h+", "Sob consulta"],
            ["Urgência", "Fora de 8×5", "+28%"],
        ],
    )
    _dx_table(doc, ["Evolução típica residual", "Porte", "Preço"], [[r[0], r[1], r[2]] for r in EVOLUCOES])
    _dx_h(doc, "6. Cronograma de implantação (Opção B)")
    _dx_table(
        doc,
        ["Semanas", "Atividade", "Responsável"],
        [
            ["1 e 2", "Infra, TLS, SMTP, backup/restore, carga", "Contratada + TI municipal"],
            ["3 e 4", "Parametrização da SME e escolas piloto", "Contratada + SME"],
            ["5 a 8", "Capacitação por papel e operação assistida", "Contratada + escolas"],
            ["9 a 12", "Go-live do recorte e relatório de implantação", "Contratada + SME"],
        ],
    )
    _dx_h(doc, "7. Níveis de serviço (Opção B)")
    _dx_table(
        doc,
        ["Severidade", "Definição", "Resposta", "Paliativo"],
        [
            ["P1", "Indisponibilidade total", "4 h úteis", "8 h úteis"],
            ["P2", "Função crítica degradada", "8 h úteis", "24 h úteis"],
            ["P3", "Falha não bloqueante", "16 h úteis", "5 dias úteis"],
            ["P4", "Evolução", "2 dias úteis", "Conforme porte"],
        ],
    )
    _dx_h(doc, "8. Propriedade intelectual")
    _dx_p(
        doc,
        "Com o pagamento integral da cessão, o contratante recebe os direitos patrimoniais de uso, modificação e exploração no âmbito de sua rede municipal. Dados pessoais permanecem de titularidade do município (LGPD). A cessão pressupõe repositório privado e licença proprietária no ato da entrega.",
    )
    _dx_h(doc, "9. Garantia")
    _dx_p(
        doc,
        "90 dias a partir da entrega do repositório, limitada a defeitos existentes na entrega e reproduzíveis. Não cobre customizações de terceiros, má operação nem requisitos novos.",
    )
    _dx_h(doc, "10. Condições gerais")
    _dx_bullets(
        doc,
        [
            f"Validade: {VALIDADE}.",
            "Preços em reais. Tributos do tomador, se devidos, serão acrescidos na NF.",
            "Contratação com a Administração observará a Lei nº 14.133/2021, quando couber.",
            "Esta proposta não gera obrigação de contratar até a assinatura do contrato.",
        ],
    )
    _dx_h(doc, "11. Aceite")
    _dx_p(doc, "Assinale a opção desejada e devolva uma via assinada:", align="left")
    _dx_table(
        doc,
        ["Opção", "Descrição", "Valor do ano 1", "Assinale"],
        [
            ["A", "Cessão do código-fonte", "R$ 520.000,00", ""],
            ["B", "Cessão + 12 meses de serviços", "R$ 540.000,00", ""],
            ["B à vista", "Opção B com 5% de desconto", "R$ 513.000,00", ""],
        ],
    )
    _dx_p(doc, "Local e data: ________________________________, ____ de _________________ de 20____.", align="center")
    _dx_p(doc, "", space_after=18)
    _dx_p(doc, "________________________________", align="center", space_after=0)
    _dx_p(doc, "Proponente", align="center", bold=True, space_after=18)
    _dx_p(doc, "________________________________", align="center", space_after=0)
    _dx_p(doc, "Destinatário (SME)", align="center", bold=True)
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    s = _styles()

    pdf_diag = OUT / "Diagnostico_Prontidao_Producao_Minima.pdf"
    pdf_prec = OUT / "Precificacao_SGE_Municipal.pdf"
    pdf_prop = OUT / "Proposta_Comercial_SGE_Municipal.pdf"
    docx_diag_p = OUT / "Diagnostico_Prontidao_Producao_Minima.docx"
    docx_prec_p = OUT / "Precificacao_SGE_Municipal.docx"
    docx_prop_p = OUT / "Proposta_Comercial_SGE_Municipal.docx"

    _build_pdf(pdf_diag, "Diagnóstico de prontidão — SGE Municipal", story_diagnostico(s))
    _build_pdf(pdf_prec, "Memorando de precificação — uso interno", story_precificacao(s))
    _build_pdf(pdf_prop, f"Proposta comercial {REF}", story_proposta(s))

    docx_diagnostico().save(docx_diag_p)
    docx_precificacao().save(docx_prec_p)
    docx_proposta().save(docx_prop_p)

    for p in (pdf_diag, pdf_prec, pdf_prop, docx_diag_p, docx_prec_p, docx_prop_p):
        print(f"{p.name}\t{p.stat().st_size}")


if __name__ == "__main__":
    main()
